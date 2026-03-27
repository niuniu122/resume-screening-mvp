from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from fastapi import HTTPException, status
from pypdf import PdfReader


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()

def extract_text_from_path(path: Path) -> tuple[str, dict]:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".docx":
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    elif suffix == ".pdf":
        reader = PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {suffix}",
        )

    normalized = normalize_text(text)
    if not normalized:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unable to extract text from file.")
    return normalized, {"suffix": suffix, "length": len(normalized)}

def extract_text_from_bytes(filename: str, content: bytes) -> tuple[str, dict]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        text = content.decode("utf-8", errors="ignore")
    elif suffix == ".docx":
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content)
            temp_path = Path(temp_file.name)
        try:
            document = Document(temp_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        finally:
            if temp_path.exists():
                temp_path.unlink()
    elif suffix == ".pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content)
            temp_path = Path(temp_file.name)
        try:
            reader = PdfReader(str(temp_path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        finally:
            if temp_path.exists():
                temp_path.unlink()
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {suffix}",
        )

    normalized = normalize_text(text)
    if not normalized:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unable to extract text from file.")
    return normalized, {"suffix": suffix, "length": len(normalized)}
